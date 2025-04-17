import os
import re
import argparse
import logging
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')


def save_docx(file_name, entries):
    """
    Save transcript entries to a DOCX file.

    Args:
        file_name (str): Path to the output DOCX file.
        entries (list): List of dictionaries containing transcript entries.
    """
    doc = Document()
    for entry in entries:
        doc.add_paragraph(entry["speaker"], style="Heading 2")
        doc.add_paragraph(entry["time"], style="Heading 3")
        doc.add_paragraph(entry["speech"], style="Normal")
    doc.save(file_name)


def save_text(file_name, entries, include_metadata=False):
    """
    Save transcript entries to a plain text file.

    Args:
        file_name (str): Path to the output text file.
        entries (list): List of dictionaries containing transcript entries.
        include_metadata (bool): If True, include speaker and timestamp metadata.
    """
    with open(file_name, 'w') as file:
        for entry in entries:
            if include_metadata:
                file.write(f'{entry["speaker"]} {entry["time"]} ')
            file.write(f'{entry["speech"]}\n')
   
    
def name_replacement(text, name_to_replace, replacement):
    """
    Replace occurrences of a given name in the text with a replacement string.

    Args:
        text (str): The input text.
        name_to_replace (str): The name to replace.
        replacement (str): The string to replace with.

    Returns:
        str: The updated text with replacements.
    """
    pattern = re.compile(r'\b' + re.escape(name_to_replace.split()[0]) + r'(?:\s+' + re.escape(" ".join(name_to_replace.split()[0][1:])) + r')?\b', re.IGNORECASE)
    return pattern.sub(replacement, text)



def process_transcript(file_path):
    """
    Process a transcript file in DOCX format and extract cleaned transcript entries.

    Args:
        file_path (str): Path to the transcript DOCX file.

    Returns:
        list: A list of dictionaries containing speaker, standardized timestamp, and speech.
    """
    # Extract PID and Name from the path
    match = re.search(r'(P0\d+)', file_path)
    pid = match.group(1) if match else None

    match = re.search(r'-\s*(.*?).docx$', file_path)
    name = match.group(1) if match else None
        
    # Open and read the document
    try:
        doc = Document(file_path)
    except Exception as e:
        logging.error(f"Error opening document {file_path}: {e}")
        return []
    
    # Step 1: Remove file's headings and footing by skipping first four and last paragraph
    paragraphs = doc.paragraphs[4:-1]

    transcript = []
    for para in paragraphs:
        para_text = para.text.strip()
        if not para_text:
            continue  # Skip empty paragraphs
        
        # Step 2: Fix spacing issues
        para_text = re.sub(r'\s+', ' ', para_text)
        para_text = re.sub(r'\n\s+', '\n', para_text)
        para_text = re.sub(r'\s+\n', '\n', para_text)
        para_text = re.sub(r'\n+', '\n', para_text)
        
        # Step 3: Extract by skipping first four and last paragraph
        pattern = r"^(.+?)\s+(\d{1,2}:\d{2}(?::\d{2})?)\s+(.*)$"
        match = re.match(pattern, para_text)
        if not match:
            logging.warning(f"Line skipped due to unmatched format: {para_text}")
            continue

        speaker = match.group(1)
        timestamp = match.group(2)
        speech = match.group(3)
        
        # TO-DO: Replace first names in the speeches
        # Replace speaker names based on PID and known interviewer name
        # If speaker is not the Interviewer, assign the PID value.
        speaker = pid if speaker != "Hoorad Abootalebi" else "Interviewer"

        # Replace the candidate's name in the speech with the PID for anonymity.
        if name:
            speech = name_replacement(speech, name, pid)
        
        # Step 5: Standardize timestamp (MM:SS or HH:MM:SS)
        parts = timestamp.split(':')
        if len(parts) == 2:
            # If it's mm:ss, assume hours = 0
            minute, second = parts
            hour = 0
        elif len(parts) == 3:
            # If it's hh:mm:ss, unpack accordingly
            hour, minute, second = parts
        else:
            raise ValueError("Timestamp format not recognized")

        standardized_timestamp = f"{int(hour):02d}:{int(minute):02d}:{int(second):02d}"
        
        # TO-DO: Find a filler word remover
        # Step 6: Remove filler words (short snetences of interviewer)
        # if speaker == "Interviewer" and len(speech) < 20:
        #     filler_words.append(speech)
        #     continue
        # speech = filler_word_remover(speech)
        
        
        # Step 7: Remove redundant or duplicated content
        
        
        # Step 8: Fix puctuation
        
        
        # Create dictionary for this transcript entry
        entry = {
            'speaker': speaker, 
            'time': standardized_timestamp, 
            'speech': speech
        }
        transcript.append(entry)
    return transcript
    

def parse_arguments():
    """
    Parse command-line arguments.

    Returns:
        argparse.Namespace: Parsed arguments including input directory and file name pattern.
    """
    parser = argparse.ArgumentParser(description="Process transcript files from a specified directory.")
    
    parser.add_argument('--input_dir', type=str, default='../data/Recordings', help="Directory containing the transcript files.")
    parser.add_argument('--start_with', type=str, default='Interview_ Social and Cultural Observations on Practices in Cybersecurity Engagement (SCOPE)', help="The pattern that transcript file names start with.")
    
    return parser.parse_args()

def process_directory(input_dir, start_with):
    """
    Walk through the given directory, process each transcript file, and save the cleaned transcript.

    Args:
        input_dir (str): The directory containing transcript files.
        start_with (str): The starting pattern of transcript file names.
    """
    for root, dirs, files in os.walk(input_dir):
        for file in files:
            if file.startswith(start_with) and file.endswith(".docx"):
                # Extract PID and Name from the path
                match = re.search(r'(P0\d+)', root)
                pid = match.group(1) if match else None

                logging.info(f'Working on participant: {pid}')
                logging.info(f'\tOpening file: {file}')
        
                transcript_file_path = os.path.join(root, file)
                clean_transcript = process_transcript(transcript_file_path)
                
                # Define the output file path for the cleaned transcript
                output_file_path = f'{root}/Interview_Transcript_{pid}.docx'

                logging.info(f'\tSaving file...')
                if not os.path.exists(output_file_path):
                    save_docx(output_file_path, clean_transcript)
                    logging.info(f'\t\tFile saved successfully: {output_file_path}')
                else:
                    logging.info(f'\t\tFile already exists: {output_file_path}')

                logging.info('*------------------------------------------------------------------------------*\n')


def main():
    """
    Main function to parse arguments and process the directory of transcript files.
    """
    args = parse_arguments()
    process_directory(args.input_dir, args.start_with)


if __name__ == "__main__":
    main()